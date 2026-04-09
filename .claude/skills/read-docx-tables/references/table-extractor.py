"""
Boilerplate: Trích xuất tất cả bảng trong .docx
— Tách struck text / surviving text / colored text per cell (không merge)
— Phân tích cấu trúc cột: header, markup, thay đổi ý nghĩa
— Section context, duplicate row detection
Chạy qua: python3 - << 'PYEOF' ... PYEOF
"""
import zipfile, re
import xml.etree.ElementTree as ET
from docx import Document

DOCX_PATH = "/path/to/file.docx"          # ← thay đường dẫn thực tế
BRD_SECTION_FILTER = None                  # None = tất cả | "INV-01" = lọc theo module name

REVIEWER_COLORS = {"EE0000", "FF0000", "C00000", "FF3300"}
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def w(tag): return f"{{{NS_W}}}{tag}"

# ─────────────────────────────────────────────────────────────────────────────
# 1. COMMENTS
# ─────────────────────────────────────────────────────────────────────────────
comments = {}
with zipfile.ZipFile(DOCX_PATH) as z:
    if "word/comments.xml" in z.namelist():
        root_c = ET.fromstring(z.read("word/comments.xml"))
        for c in root_c.findall(f"{{{NS_W}}}comment"):
            cid    = c.get(w("id"))
            author = c.get(w("author"), "")
            texts  = [t.text or "" for t in c.iter(w("t"))]
            comments[cid] = {"author": author, "text": " ".join(texts).strip()}

# ─────────────────────────────────────────────────────────────────────────────
# 2. MARKUP HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def get_comment_ids(para):
    xml_str = para._element.xml
    if "commentRangeStart" not in xml_str:
        return []
    return re.findall(r'commentRangeStart[^>]+w:id="(\d+)"', xml_str)

def get_cell_data(cell):
    """
    Trả về dict markup đầy đủ của một cell — tách rõ từng đoạn theo loại.

    segments: list of (type, text) theo thứ tự xuất hiện trong cell.
      type = "normal" | "struck" | "colored" | "highlight"
    surviving_text: text còn lại sau khi bỏ struck (= nội dung mới)
    struck_text:    text bị gạch bỏ (= nội dung cũ cần xóa)
    """
    segments    = []   # [(type, text), ...]  — ordered, per-run
    comment_ids = []

    for para in cell.paragraphs:
        for run in para.runs:
            t = run.text or ""
            if not t:
                continue

            rpr = run._element.find(w("rPr"))
            color_val = ""
            if rpr is not None:
                col_el = rpr.find(w("color"))
                if col_el is not None:
                    color_val = col_el.get(w("val"), "").upper()

            is_struck    = bool(run.font.strike)
            is_colored   = color_val in REVIEWER_COLORS
            is_highlight = run.font.highlight_color is not None

            if is_struck:
                segments.append(("struck", t))
            elif is_colored:
                segments.append(("colored", t))
            elif is_highlight:
                segments.append(("highlight", t))
            else:
                segments.append(("normal", t))

        comment_ids.extend(get_comment_ids(para))

    # Derived aggregates
    surviving = "".join(t for tp, t in segments if tp != "struck").strip()
    struck    = "".join(t for tp, t in segments if tp == "struck").strip()
    colored   = "".join(t for tp, t in segments if tp == "colored").strip()
    highlight = "".join(t for tp, t in segments if tp == "highlight").strip()

    return {
        "segments":       segments,
        "full_text":      cell.text.strip(),   # toàn bộ text (kể cả struck)
        "surviving_text": surviving,            # text còn lại = nội dung mới
        "struck_text":    struck,               # text bị gạch = nội dung cũ
        "colored_text":   colored,              # reviewer thêm mới
        "highlight_text": highlight,            # highlight để chú ý
        "comment_ids":    comment_ids,
        "has_markup":     bool(struck or colored or highlight or comment_ids),
    }

def render_cell(c, max_len=55):
    """
    Render cell theo format rõ ràng:
      Không markup → surviving_text
      Có struck     → ~~struck~~ surviving
      Có colored    → surviving [+colored]
      Kết hợp       → ~~struck~~ surviving [+colored]
    Thứ tự theo segments để giữ đúng trình tự trong tài liệu.
    """
    if not c["has_markup"] and not c["comment_ids"]:
        return c["full_text"][:max_len]

    parts = []
    for tp, t in c["segments"]:
        t_s = t.strip()
        if not t_s:
            continue
        if tp == "struck":
            parts.append(f"~~{t_s[:30]}~~")
        elif tp == "colored":
            parts.append(f"[+{t_s[:30]}]")
        elif tp == "highlight":
            parts.append(f"[HL:{t_s[:30]}]")
        else:
            parts.append(t_s[:30])

    cell_str = " ".join(parts) if parts else c["full_text"][:max_len]

    # Gắn comment
    for cid in c["comment_ids"]:
        cinfo = comments.get(cid, {})
        cell_str += f" 💬C{cid}:{cinfo.get('text','')[:40]}"

    return cell_str

# ─────────────────────────────────────────────────────────────────────────────
# 3. DUYỆT TÀI LIỆU — giữ thứ tự, gắn section heading
# ─────────────────────────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)
para_map  = {p._element: p for p in doc.paragraphs}
table_map = {t._element: t for t in doc.tables}

def is_heading(para):
    return para.style.name.lower().startswith("heading")

def heading_level(para):
    try:
        return int(para.style.name.lower().replace("heading", "").strip())
    except:
        return 0

current_module  = "【Đầu tài liệu】"   # Heading 1 — dùng để filter
current_section = "【Đầu tài liệu】"   # Heading gần nhất — dùng để hiển thị

tables_out = []
t_idx_global = 0

for child in doc.element.body.iterchildren():
    tag = child.tag
    if tag == w("p"):
        para = para_map.get(child)
        if para and is_heading(para) and para.text.strip():
            level = heading_level(para)
            current_section = para.text.strip()
            if level == 1:
                current_module = para.text.strip()

    elif tag == w("tbl"):
        tbl = table_map.get(child)
        if tbl is None:
            continue
        t_idx_global += 1

        if BRD_SECTION_FILTER and BRD_SECTION_FILTER.lower() not in current_module.lower():
            continue

        rows_data = []
        for row in tbl.rows:
            rows_data.append([get_cell_data(cell) for cell in row.cells])

        if not rows_data:
            continue

        header_row  = rows_data[0]
        header_vals = [c["surviving_text"] or c["full_text"] for c in header_row]
        # Phát hiện markup trong header
        header_markup = [(i, c) for i, c in enumerate(header_row) if c["has_markup"]]

        tables_out.append({
            "t_idx":          t_idx_global,
            "module":         current_module,
            "section":        current_section,
            "header":         header_vals,
            "header_row":     header_row,
            "header_markup":  header_markup,   # cột có markup → có thể thay đổi
            "rows":           rows_data,
        })

# ─────────────────────────────────────────────────────────────────────────────
# 4. OUTPUT
# ─────────────────────────────────────────────────────────────────────────────
meaningful = [t for t in tables_out if len(t["rows"]) >= 2 and len(t["header"]) >= 2]

print(f"📊 Tổng: {t_idx_global} bảng trong docx | {len(meaningful)} bảng có nội dung (≥2 cột, ≥2 row)")
print("=" * 80)

for tbl in meaningful:
    header       = tbl["header"]
    header_row   = tbl["header_row"]
    rows         = tbl["rows"]
    header_mkup  = tbl["header_markup"]

    n_markup = sum(1 for row in rows for cell in row if cell["has_markup"])

    print(f"\n[T{tbl['t_idx']}] Module  : {tbl['module'][:65]}")
    print(f"  Section : {tbl['section'][:65]}")
    print(f"  Rows: {len(rows)} | Cells có markup: {n_markup}")

    # ── Phân tích cột (header) ──────────────────────────────────────────────
    print(f"\n  📋 CỘT ({len(header)}):")
    for i, col in enumerate(header):
        hc   = header_row[i]
        flag = ""
        if hc["struck_text"]:    flag += f" [TÊN CŨ: ~~{hc['struck_text'][:25]}~~]"
        if hc["colored_text"]:   flag += f" [TÊN MỚI: +{hc['colored_text'][:25]}]"
        if hc["highlight_text"]: flag += f" [HL]"
        if hc["comment_ids"]:
            for cid in hc["comment_ids"]:
                flag += f" 💬C{cid}:{comments.get(cid,{}).get('text','')[:30]}"
        print(f"    [{i}] {col}{flag}")

    if header_mkup:
        print(f"  ⚠️  {len(header_mkup)} cột có markup trong header → kiểm tra thay đổi tên/ý nghĩa cột")

    print()

    # ── Nội dung rows ───────────────────────────────────────────────────────
    seen = set()
    for r_idx, row in enumerate(rows):
        row_key = tuple(c["full_text"][:30] for c in row)
        if r_idx > 0 and row_key == tuple(rows[r_idx-1][i]["full_text"][:30] for i in range(len(row))):
            continue
        if row_key in seen:
            continue
        seen.add(row_key)

        cells_display = [render_cell(c) for c in row]
        prefix = "  HDR│" if r_idx == 0 else f"  R{r_idx:02d}│"
        print(prefix + " | ".join(cells_display))

    print()
    print("  " + "─" * 76)

# ─────────────────────────────────────────────────────────────────────────────
# 5. PHÂN TÍCH CỘT TỔNG HỢP — bảng nào thay đổi cột
# ─────────────────────────────────────────────────────────────────────────────
print("\n" + "=" * 80)
print("\n📋 PHÂN TÍCH CỘT — BẢNG CÓ THAY ĐỔI CẤU TRÚC:\n")

for tbl in meaningful:
    header_mkup = tbl["header_markup"]
    header      = tbl["header"]
    rows        = tbl["rows"]

    # Phát hiện cột mới / cột bị xóa / cột đổi tên
    col_changes = []
    for i, hc in tbl["header_markup"]:
        old_name = hc["struck_text"]
        new_name = hc["colored_text"] or (hc["surviving_text"] if hc["struck_text"] else "")
        if old_name and new_name:
            col_changes.append(f"  ĐỔI TÊN cột [{i}]: ~~{old_name}~~ → {new_name}")
        elif old_name and not new_name:
            col_changes.append(f"  XÓA cột [{i}]: ~~{old_name}~~")
        elif not old_name and hc["colored_text"]:
            col_changes.append(f"  THÊM nội dung cột [{i}]: +{hc['colored_text']}")

    # Phát hiện cột có nhiều highlighted rows → thay đổi ý nghĩa
    col_hl_count = [0] * len(header)
    for row in rows[1:]:  # bỏ header row
        for j, cell in enumerate(row):
            if j < len(col_hl_count) and cell["highlight_text"]:
                col_hl_count[j] += 1

    hl_cols = [(j, header[j], cnt) for j, cnt in enumerate(col_hl_count) if cnt >= 2]

    if col_changes or hl_cols or header_mkup:
        print(f"  [T{tbl['t_idx']}] {tbl['section'][:60]}")
        print(f"       Cột hiện tại ({len(header)}): {' | '.join(h[:20] for h in header)}")
        for ch in col_changes:
            print(ch)
        if hl_cols:
            for j, col_name, cnt in hl_cols:
                print(f"  ⚠️  Cột [{j}] '{col_name}': {cnt} rows được highlight → có thể thay đổi ý nghĩa/nội dung")
        print()

# ─────────────────────────────────────────────────────────────────────────────
# 6. TÓM TẮT — bảng nào cần so sánh với BRD
# ─────────────────────────────────────────────────────────────────────────────
print("=" * 80)
print("\n🔍 BẢNG CẦN SO SÁNH VỚI BRD:\n")
for tbl in meaningful:
    rows    = tbl["rows"]
    n_markup = sum(1 for row in rows for cell in row if cell["has_markup"])
    n_comments = sum(len(cell["comment_ids"]) for row in rows for cell in row)
    if n_markup > 0 or len(rows) > 5:
        col_note = ""
        if tbl["header_markup"]:
            col_note = f" | ⚠️ {len(tbl['header_markup'])} cột thay đổi"
        print(f"  [T{tbl['t_idx']}] {tbl['section'][:65]}")
        print(f"       {len(rows)} rows | {len(tbl['header'])} cột{col_note} | {n_markup} cells có markup | {n_comments} comments")
        print(f"       Cột: {' | '.join(h[:20] for h in tbl['header'])}")
        print()

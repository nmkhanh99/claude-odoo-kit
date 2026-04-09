"""
Boilerplate: Khảo sát toàn bộ loại markup trong file .docx
Mục đích: Chạy TRƯỚC extract-docx-comments để biết file có gì, cấu hình đúng tham số

Output:
  - Bảng tổng hợp: loại markup → số lượng → có nên trích xuất không
  - Sample cho mỗi loại (để confirm trước khi extract)
  - Gợi ý cấu hình REVIEWER_COLORS và HIGHLIGHT_FILTER

Chạy qua: python3 - << 'PYEOF' ... PYEOF
"""
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from collections import Counter
import re

DOCX_PATH = "/path/to/file.docx"   # ← thay đường dẫn thực tế

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def w(tag): return f"{{{NS_W}}}{tag}"

# ─────────────────────────────────────────────────────────────────────────────
# 0. Liệt kê các file XML trong docx (để biết file có phần nào)
# ─────────────────────────────────────────────────────────────────────────────
with zipfile.ZipFile(DOCX_PATH) as z:
    zip_files = z.namelist()
    has_comments      = "word/comments.xml" in zip_files
    has_people        = "word/people.xml" in zip_files
    has_footnotes     = "word/footnotes.xml" in zip_files
    has_endnotes      = "word/endnotes.xml" in zip_files
    doc_xml           = z.read("word/document.xml").decode("utf-8")
    comments_xml      = z.read("word/comments.xml").decode("utf-8") if has_comments else ""

print("=" * 70)
print("📂 THÀNH PHẦN FILE DOCX")
print("=" * 70)
print(f"  comments.xml   : {'✅' if has_comments else '❌ (không có comment)'}")
print(f"  people.xml     : {'✅ (có thông tin reviewer)' if has_people else '❌'}")
print(f"  footnotes.xml  : {'✅' if has_footnotes else '❌'}")
print(f"  endnotes.xml   : {'✅' if has_endnotes else '❌'}")
print()

# ─────────────────────────────────────────────────────────────────────────────
# 1. Comments — đếm và lấy danh sách author
# ─────────────────────────────────────────────────────────────────────────────
comment_count   = 0
comment_authors = Counter()

if has_comments:
    root_c = ET.fromstring(comments_xml)
    for c in root_c.findall(f"{{{NS_W}}}comment"):
        comment_count += 1
        author = c.get(w("author"), "unknown")
        comment_authors[author] += 1

# ─────────────────────────────────────────────────────────────────────────────
# 2. Scan raw XML cho các loại markup
# ─────────────────────────────────────────────────────────────────────────────
ins_count          = len(re.findall(r'<w:ins\b', doc_xml))
del_count          = len(re.findall(r'<w:del\b', doc_xml))
strike_count       = len(re.findall(r'<w:strike(?:\s[^/]*)?\s*/>', doc_xml))
highlight_count    = len(re.findall(r'<w:highlight\b', doc_xml))
comment_ref_count  = len(re.findall(r'<w:commentRangeStart\b', doc_xml))

# Color values (non-black, non-auto)
color_vals = Counter(re.findall(r'<w:color\s+w:val="([A-Fa-f0-9]{6})"', doc_xml))
color_vals.pop("000000", None)  # loại bỏ đen
# Shading fill values
shd_fills = Counter(re.findall(r'<w:shd\b[^>]*w:fill="([A-Fa-f0-9]{6})"', doc_xml))
shd_fills.pop("000000", None); shd_fills.pop("FFFFFF", None); shd_fills.pop("auto", None)
# Highlight values
hl_vals = Counter(re.findall(r'<w:highlight\s+w:val="([^"]+)"', doc_xml))

# ─────────────────────────────────────────────────────────────────────────────
# 3. Scan python-docx cho sample content
# ─────────────────────────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)

strike_samples    = []
highlight_samples = []
color_samples     = {}   # color_hex → [(text, context)]

def scan_para(para, source="paragraph"):
    for run in para.runs:
        # Strike
        if run.font.strike and run.text.strip() and len(strike_samples) < 3:
            strike_samples.append({
                "source": source,
                "text": run.text[:80],
                "context": para.text[:80],
            })
        # Highlight
        if run.font.highlight_color is not None and run.text.strip() and len(highlight_samples) < 3:
            highlight_samples.append({
                "source": source,
                "color": str(run.font.highlight_color).split(".")[-1],
                "text": run.text[:60],
                "context": para.text[:80],
            })
        # Color
        rpr = run._element.find(w("rPr"))
        if rpr is not None:
            col_el = rpr.find(w("color"))
            if col_el is not None:
                val = col_el.get(w("val"), "").upper()
                if val and val not in ("000000", "AUTO") and run.text.strip():
                    if val not in color_samples:
                        color_samples[val] = []
                    if len(color_samples[val]) < 2:
                        color_samples[val].append({"source": source, "text": run.text[:60], "context": para.text[:60]})

for para in doc.paragraphs:
    scan_para(para, "paragraph")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                scan_para(para, "table")

# ─────────────────────────────────────────────────────────────────────────────
# 4. OUTPUT
# ─────────────────────────────────────────────────────────────────────────────
print("=" * 70)
print("📊 TỔNG HỢP MARKUP")
print("=" * 70)
print(f"  💬 Comments           : {comment_count:4d}  ({'✅ có' if comment_count else '❌ không'})")
print(f"  🔗 Comment refs map    : {comment_ref_count:4d}  (commentRangeStart trong body)")
print(f"  ✂️  Strikethrough       : {strike_count:4d}  (w:strike trong rPr)")
print(f"  🔄 Tracked INSERT      : {ins_count:4d}  (w:ins – Word revision tracking)")
print(f"  🔄 Tracked DELETE      : {del_count:4d}  (w:del – Word revision tracking)")
print(f"  🟡 Highlights          : {highlight_count:4d}  (w:highlight)")
print(f"  🎨 Colored text runs   : {sum(color_vals.values()):4d}  ({len(color_vals)} màu khác nhau)")
print(f"  🟦 Shading (background): {sum(shd_fills.values()):4d}  ({len(shd_fills)} màu khác nhau)")

# ── Comment authors
if comment_authors:
    print(f"\n👥 REVIEWER ({len(comment_authors)} người):")
    for author, cnt in comment_authors.most_common():
        print(f"     {cnt:3d} comments — {author}")

# ── Highlight values
if hl_vals:
    print(f"\n🟡 HIGHLIGHT COLORS:")
    for val, cnt in hl_vals.most_common():
        print(f"     {cnt:4d}x  {val}")

# ── Color values
if color_vals:
    print(f"\n🎨 TEXT COLORS (top 15, non-black):")
    for val, cnt in color_vals.most_common(15):
        samples = color_samples.get(val.upper(), [])
        sample_text = f'— vd: "{samples[0]["text"][:40]}"' if samples else ""
        print(f"     #{val.upper():6s}  {cnt:4d}x  {sample_text}")

# ── Shading
if shd_fills:
    print(f"\n🟦 SHADING FILLS (top 10):")
    for val, cnt in shd_fills.most_common(10):
        print(f"     #{val.upper():6s}  {cnt:4d}x")

# ── Samples
if strike_samples:
    print(f"\n✂️  STRIKE SAMPLES:")
    for s in strike_samples:
        print(f"  [{s['source']}] Gạch bỏ: \"{s['text']}\"")
        print(f"    Ngữ cảnh: \"{s['context']}\"")

if highlight_samples:
    print(f"\n🟡 HIGHLIGHT SAMPLES:")
    for h in highlight_samples:
        print(f"  [{h['source']}] [{h['color']}] \"{h['text']}\"")
        print(f"    Ngữ cảnh: \"{h['context']}\"")

# ─────────────────────────────────────────────────────────────────────────────
# 5. Gợi ý cấu hình cho extract-docx-comments
# ─────────────────────────────────────────────────────────────────────────────
print()
print("=" * 70)
print("💡 GỢI Ý CẤU HÌNH CHO extract-docx-comments")
print("=" * 70)

# Phát hiện màu reviewer (thường là đỏ, cam, tím – không phải branding)
reviewer_candidates = []
branding_threshold  = 100   # màu xuất hiện > 100 lần thường là branding
for val, cnt in color_vals.most_common():
    upper = val.upper()
    # Bỏ qua màu branding phổ biến (trắng trên header, xám, v.v.)
    if cnt < branding_threshold:
        # Phân loại theo tone
        r_hex = int(upper[0:2], 16)
        g_hex = int(upper[2:4], 16)
        b_hex = int(upper[4:6], 16)
        # Đỏ: r cao, g và b thấp
        if r_hex > 150 and g_hex < 100 and b_hex < 100:
            reviewer_candidates.append(f'"{upper}"  # đỏ ({cnt}x)')
        # Tím/cam: r cao, g thấp-trung
        elif r_hex > 150 and g_hex < 150 and b_hex < 80:
            reviewer_candidates.append(f'"{upper}"  # cam/tím ({cnt}x)')

if reviewer_candidates:
    print(f"\n  REVIEWER_COLORS gợi ý (màu chữ reviewer):")
    for c in reviewer_candidates[:8]:
        print(f"    {c}")
else:
    print(f"\n  REVIEWER_COLORS: Không phát hiện rõ. Xem bảng 🎨 ở trên và chọn màu phù hợp.")

# Highlight filter
if hl_vals:
    dominant_hl = hl_vals.most_common(1)[0]
    print(f"\n  HIGHLIGHT_FILTER: Chủ yếu là '{dominant_hl[0]}' ({dominant_hl[1]}x)")
    print(f"    → Dùng mặc định (None = tất cả) hoặc lọc theo màu nếu cần")

# Tracked changes
if ins_count > 0 or del_count > 0:
    print(f"\n  ✅ File có Tracked Changes (INSERT: {ins_count}, DELETE: {del_count})")
    print(f"     → extract-docx-comments sẽ tự động trích xuất")
else:
    print(f"\n  ℹ️  Không có Tracked Changes (w:ins/w:del)")
    print(f"     → Nếu có thay đổi văn bản, reviewer dùng strike thay vì Word revision")

# Strike vs Tracked changes
if strike_count > 0 and del_count == 0:
    print(f"\n  ⚠️  {strike_count} strike nhưng 0 w:del → Reviewer gạch bỏ thủ công (không dùng Track Changes)")
    print(f"     → Cần đọc cả strikethrough lẫn comment để hiểu đầy đủ thay đổi")

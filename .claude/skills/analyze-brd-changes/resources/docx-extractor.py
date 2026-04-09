"""
Boilerplate trích xuất Comments + Tracked Changes từ .docx
Chạy qua: python3 - << 'PYEOF' ... PYEOF
"""
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

DOCX_PATH = "/path/to/feedback.docx"

# ── 1. Trích xuất Comments từ word/comments.xml ────────────────────────────
NS = {
    "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}

comments = {}  # {comment_id: {"author": ..., "date": ..., "text": ...}}

with zipfile.ZipFile(DOCX_PATH) as z:
    if "word/comments.xml" in z.namelist():
        root = ET.fromstring(z.read("word/comments.xml"))
        for c in root.findall("w:comment", NS):
            cid    = c.get(f'{{{NS["w"]}}}id')
            author = c.get(f'{{{NS["w"]}}}author', "")
            date   = c.get(f'{{{NS["w"]}}}date', "")[:10]  # YYYY-MM-DD
            texts  = [r.text or "" for r in c.iter(f'{{{NS["w"]}}}t')]
            comments[cid] = {
                "author": author,
                "date":   date,
                "text":   " ".join(texts).strip(),
            }

print(f"Tìm thấy {len(comments)} comments")
for cid, info in sorted(comments.items(), key=lambda x: int(x[0])):
    print(f"  C{cid} – {info['author']} – {info['date']}: {info['text'][:80]}")

# ── 2. Trích xuất Tracked Changes (gạch bỏ) ───────────────────────────────
doc = Document(DOCX_PATH)

strikethrough_texts = []

def extract_strikes(paragraphs):
    for para in paragraphs:
        strike_runs = [r.text for r in para.runs if r.font.strike and r.text.strip()]
        if strike_runs:
            strikethrough_texts.append({
                "location": para.style.name,
                "text":     " ".join(strike_runs),
            })

# Paragraphs
extract_strikes(doc.paragraphs)

# Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            extract_strikes(cell.paragraphs)

print(f"\nTìm thấy {len(strikethrough_texts)} đoạn gạch bỏ:")
for item in strikethrough_texts:
    print(f"  [{item['location']}] {item['text'][:80]}")

# ── 3. Map comment vào đoạn văn (comment reference) ───────────────────────
# Comment references nằm trong paragraph XML dưới dạng <w:commentRangeStart w:id="X"/>
# Đọc raw XML để map:

comment_refs = {}  # {comment_id: paragraph_text}

for para in doc.paragraphs:
    xml_str = para._element.xml
    if "commentRangeStart" in xml_str:
        import re
        ids = re.findall(r'commentRangeStart[^>]+w:id="(\d+)"', xml_str)
        for cid in ids:
            comment_refs[cid] = para.text[:120]

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                xml_str = para._element.xml
                if "commentRangeStart" in xml_str:
                    import re
                    ids = re.findall(r'commentRangeStart[^>]+w:id="(\d+)"', xml_str)
                    for cid in ids:
                        comment_refs[cid] = para.text[:120]

print(f"\nMap được {len(comment_refs)} comment → đoạn văn:")
for cid, para_text in sorted(comment_refs.items(), key=lambda x: int(x[0])):
    info = comments.get(cid, {})
    print(f"  C{cid} [{info.get('author','')}]: \"{para_text[:60]}\"")
    print(f"       → {info.get('text','')[:80]}")
